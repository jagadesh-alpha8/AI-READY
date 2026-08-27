"""Renders a Report's `report_data` into downloadable PDF/DOCX bytes.

fpdf2's built-in core fonts (Helvetica) only support Latin-1 -- rather than
bundling a Unicode TrueType font (shipping a font file, or depending on an
OS-specific system font path, neither of which is portable across the
dev/CI/production environments this runs in), dynamic text is passed
through `_latin1` first: any character outside Latin-1 is replaced with '?'
rather than crashing the render. python-docx has no such restriction
(docx/XML is UTF-8), so DOCX rendering uses the real text as-is.
"""
import io

from docx import Document as DocxDocument
from fpdf import FPDF


def _latin1(text):
    return str(text).encode('latin-1', 'replace').decode('latin-1')


def _line(pdf, height, text):
    """multi_cell that always leaves the cursor back at the left margin on
    the next line -- fpdf2 otherwise leaves it wherever the text ended,
    which can leave near-zero horizontal room for the following call."""
    pdf.multi_cell(0, height, text, new_x='LMARGIN', new_y='NEXT')


# ---------------------------------------------------------------- PDF -----

def render_pdf_bytes(report):
    data = report.report_data
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font('Helvetica', 'B', 18)
    _line(pdf, 10, _latin1(f"AIOS Discovery Report - {data['institution']['name']}"))
    pdf.set_font('Helvetica', '', 10)
    _line(pdf, 6, _latin1(
        f"Sprint {data['sprint']['sprint_code']} | Version {report.version} | Generated {data['generated_at']}",
    ))
    pdf.set_font('Helvetica', 'B', 10)
    pdf.set_text_color(150, 90, 0)
    _line(pdf, 6, _latin1(f"Baseline status: {data['baseline']['status'].upper()}"))
    pdf.set_text_color(0, 0, 0)
    pdf.ln(4)

    _section(pdf, '1. Executive Summary', [data['executive_summary']])

    _heading(pdf, '2. Overall CRI & Confidence')
    pdf.set_font('Helvetica', '', 11)
    _line(pdf, 7, _latin1(f"Overall CRI: {data['overall_cri']:.1f} / 100"))
    _line(pdf, 7, _latin1(f"Confidence: {data['confidence_score']:.0%}"))
    pdf.ln(2)

    _section(pdf, '3. Eight Pillar Scorecards', [
        f"{p['label']}: {p['raw_score']:.1f}/100 ({p['status']}) - confidence {p['confidence_score']:.0%}"
        for p in data['pillar_scorecards']
    ] or ['No pillar scores yet.'])

    _section(pdf, '4. Strengths', [
        f"{s['label']} - {s['raw_score']:.1f}/100" for s in data['strengths']
    ] or ['None yet identified.'])

    _section(pdf, '5. Areas for Improvement', [
        f"{s['label']} - {s['raw_score']:.1f}/100 ({s['status']})" for s in data['areas_for_improvement']
    ] or ['None.'])

    _section(pdf, '6. Missing Data Appendix', [
        f"[{g['priority']}] {g['pillar_label']}: {g['title']}" for g in data['missing_data_appendix']
    ] or ['No open data gaps.'])

    _section(pdf, '7. Recommendations', [
        f"[{r['priority']}] {r['title']} (+{r['expected_cri_lift']:.1f} CRI lift)"
        for r in data['recommendations']
    ] or ['No recommendations generated yet.'])

    _heading(pdf, '8. 90-Day Action Plan')
    _action_plan_body(pdf, data['ninety_day_action_plan'])

    _heading(pdf, '9. 12-Month Roadmap')
    _action_plan_body(pdf, data['twelve_month_roadmap'])

    _section(pdf, '10. How InGage Can Help', [
        f"{o['offering']} ({o['recommendation_count']} recommendation(s))"
        for o in data['how_ingage_can_help']
    ] or ['No support offerings linked yet.'])

    return bytes(pdf.output())


def _heading(pdf, text):
    pdf.set_font('Helvetica', 'B', 13)
    pdf.set_text_color(20, 60, 120)
    _line(pdf, 8, _latin1(text))
    pdf.set_text_color(0, 0, 0)


def _bullets(pdf, lines):
    pdf.set_font('Helvetica', '', 9)
    for line in lines:
        _line(pdf, 5.5, _latin1(f'- {line}'))
    pdf.ln(2)


def _section(pdf, title, lines):
    _heading(pdf, title)
    _bullets(pdf, lines)


def _action_plan_body(pdf, buckets):
    if not buckets:
        _bullets(pdf, ['No items identified for this horizon.'])
        return
    for bucket in buckets:
        pdf.set_font('Helvetica', 'B', 10)
        _line(pdf, 6, _latin1(bucket['timeline']))
        _bullets(pdf, [item['title'] for item in bucket['items']])


# --------------------------------------------------------------- DOCX -----

def render_docx_bytes(report):
    data = report.report_data
    doc = DocxDocument()

    doc.add_heading(f"AIOS Discovery Report — {data['institution']['name']}", level=0)
    meta = doc.add_paragraph()
    meta.add_run(
        f"Sprint {data['sprint']['sprint_code']} · Version {report.version} · Generated {data['generated_at']}",
    ).italic = True
    baseline_run = doc.add_paragraph().add_run(f"Baseline status: {data['baseline']['status'].upper()}")
    baseline_run.bold = True

    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(data['executive_summary'])

    doc.add_heading('2. Overall CRI & Confidence', level=1)
    doc.add_paragraph(f"Overall CRI: {data['overall_cri']:.1f} / 100")
    doc.add_paragraph(f"Confidence: {data['confidence_score']:.0%}")

    doc.add_heading('3. Eight Pillar Scorecards', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    for cell, label in zip(table.rows[0].cells, ['Pillar', 'Score', 'Status', 'Confidence', 'Evidence']):
        cell.text = label
    for p in data['pillar_scorecards']:
        row = table.add_row().cells
        row[0].text = p['label']
        row[1].text = f"{p['raw_score']:.1f}/100"
        row[2].text = p['status']
        row[3].text = f"{p['confidence_score']:.0%}"
        row[4].text = str(p['evidence_count'])

    _docx_bullets(doc, '4. Strengths', [f"{s['label']} — {s['raw_score']:.1f}/100" for s in data['strengths']])
    _docx_bullets(
        doc, '5. Areas for Improvement',
        [f"{s['label']} — {s['raw_score']:.1f}/100 ({s['status']})" for s in data['areas_for_improvement']],
    )
    _docx_bullets(
        doc, '6. Missing Data Appendix',
        [f"[{g['priority']}] {g['pillar_label']}: {g['title']}" for g in data['missing_data_appendix']],
    )
    _docx_bullets(
        doc, '7. Recommendations',
        [
            f"[{r['priority']}] {r['title']} (+{r['expected_cri_lift']:.1f} CRI lift)"
            for r in data['recommendations']
        ],
    )

    doc.add_heading('8. 90-Day Action Plan', level=1)
    _docx_action_plan(doc, data['ninety_day_action_plan'])

    doc.add_heading('9. 12-Month Roadmap', level=1)
    _docx_action_plan(doc, data['twelve_month_roadmap'])

    _docx_bullets(
        doc, '10. How InGage Can Help',
        [f"{o['offering']} ({o['recommendation_count']} recommendation(s))" for o in data['how_ingage_can_help']],
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _docx_bullets(doc, heading, lines):
    doc.add_heading(heading, level=1)
    if not lines:
        doc.add_paragraph('None.')
        return
    for line in lines:
        doc.add_paragraph(line, style='List Bullet')


def _docx_action_plan(doc, buckets):
    if not buckets:
        doc.add_paragraph('No items identified for this horizon.')
        return
    for bucket in buckets:
        doc.add_paragraph().add_run(bucket['timeline']).bold = True
        for item in bucket['items']:
            doc.add_paragraph(item['title'], style='List Bullet')
